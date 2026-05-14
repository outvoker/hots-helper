"""Convert REPORT-style Markdown into a DOCX document.

Hand-rolled to keep the dep surface tiny and avoid pandoc/LibreOffice. Handles
the subset of Markdown actually used in our reports:

- ATX headings (``#`` to ``####``)
- Paragraphs with inline ``**bold**``, ``*italic*``, ``_italic_`` and ``code``
- GitHub-style pipe tables (``|---|---:|`` headers, with right/left/center alignment)
- Blockquotes (``> `` prefix)
- Bullet lists (``- ``)
- Horizontal rules (``---``)

Usage:
    uv run python scripts/md_to_docx.py REPORT.zh.md REPORT.zh.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


# --- inline parsing ---------------------------------------------------------

# Order matters: longer markers first.
_INLINE_RE = re.compile(
    r"(\*\*([^*]+)\*\*"      # bold
    r"|\*([^*]+)\*"            # italic with *
    r"|_([^_]+)_"              # italic with _
    r"|`([^`]+)`)"             # code
)


def _add_runs(paragraph, text: str) -> None:
    """Append `text` to `paragraph`, parsing **bold**, *italic*, `code`."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        bold, italic_a, italic_u, code = m.group(2), m.group(3), m.group(4), m.group(5)
        if bold is not None:
            run = paragraph.add_run(bold)
            run.bold = True
        elif italic_a is not None:
            run = paragraph.add_run(italic_a)
            run.italic = True
        elif italic_u is not None:
            run = paragraph.add_run(italic_u)
            run.italic = True
        elif code is not None:
            run = paragraph.add_run(code)
            run.font.name = "Menlo"
            # East-Asian font fallback for CJK in code spans:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "Menlo")
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# --- table parser -----------------------------------------------------------


def _is_table_separator(line: str) -> bool:
    """Match a line like ``|---|---:|:---:|---|`` used as table delim."""
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def _parse_alignments(separator: str) -> list[str]:
    out = []
    for c in [c.strip() for c in separator.strip("|").split("|")]:
        if c.startswith(":") and c.endswith(":"):
            out.append("center")
        elif c.endswith(":"):
            out.append("right")
        else:
            out.append("left")
    return out


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [cell.strip() for cell in s.split("|")]


def _emit_table(doc, header_row: list[str], align: list[str], rows: list[list[str]]):
    cols = max(len(header_row), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for i in range(cols):
        cell = table.rows[0].cells[i]
        cell.text = ""
        text = header_row[i] if i < len(header_row) else ""
        p = cell.paragraphs[0]
        _set_cell_alignment(p, align[i] if i < len(align) else "left")
        _add_runs(p, text)
        for run in p.runs:
            run.bold = True

    # Body
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            _set_cell_alignment(p, align[c_idx] if c_idx < len(align) else "left")
            _add_runs(p, text)


def _set_cell_alignment(paragraph, align: str) -> None:
    if align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


# --- main converter ---------------------------------------------------------


def _set_document_default_font(doc: Document, font: str = "PingFang SC") -> None:
    """Force a CJK-friendly default font so 中文 renders correctly on Windows/Mac."""
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    # ascii / hAnsi / eastAsia all need to be set; if PingFang SC is missing
    # on Windows, Office falls back to Microsoft YaHei or SimSun.
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)


def _add_heading(doc, level: int, text: str) -> None:
    p = doc.add_heading(level=level)
    _add_runs(p, text)


def _consume_table(lines: list[str], i: int) -> tuple[list[str], list[str], list[list[str]], int]:
    header = _split_row(lines[i])
    align = _parse_alignments(lines[i + 1])
    j = i + 2
    rows = []
    while j < len(lines) and lines[j].strip().startswith("|"):
        rows.append(_split_row(lines[j]))
        j += 1
    return header, align, rows, j


def convert(md_path: Path, docx_path: Path) -> None:
    md = md_path.read_text("utf-8")
    lines = md.splitlines()

    doc = Document()
    _set_document_default_font(doc, "PingFang SC")

    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # Empty line
        if not s:
            i += 1
            continue

        # Horizontal rule
        if s in {"---", "***", "___"}:
            doc.add_paragraph().add_run().add_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            _add_heading(doc, min(level, 4), m.group(2))
            i += 1
            continue

        # Tables
        if (
            s.startswith("|")
            and i + 1 < len(lines)
            and _is_table_separator(lines[i + 1])
        ):
            header, align, rows, ni = _consume_table(lines, i)
            _emit_table(doc, header, align, rows)
            i = ni
            continue

        # Blockquote
        if s.startswith("> "):
            quote_lines: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith("> "):
                quote_lines.append(lines[i].lstrip()[2:])
                i += 1
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, " ".join(quote_lines))
            continue

        # Bullet list
        if s.startswith("- ") or s.startswith("* "):
            while i < len(lines) and lines[i].lstrip().startswith(("- ", "* ")):
                content = lines[i].lstrip()[2:]
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, content)
                i += 1
            continue

        # Plain paragraph (collapse consecutive non-empty lines)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].lstrip().startswith(("#", "- ", "* ", ">", "|"))
        ):
            para_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        _add_runs(p, " ".join(line.strip() for line in para_lines))

    doc.save(str(docx_path))


def main() -> int:
    if len(sys.argv) < 3:
        print("usage: md_to_docx.py <input.md> <output.docx>")
        return 2
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"wrote {sys.argv[2]}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
